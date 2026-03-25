import docx
import os

def extract_docx_content(file_path):
    if not os.path.exists(file_path):
        print(f"File {file_path} not found.")
        return

    doc = docx.Document(file_path)
    
    print(f"--- CONTENT OF {file_path} ---")
    for para in doc.paragraphs:
        if para.text.strip():
            print(f"P: {para.text}")
    
    for table in doc.tables:
        print("\n--- TABLE ---")
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            print(f"TR: {' | '.join(row_text)}")

if __name__ == "__main__":
    extract_docx_content("DS Record - Initial Pages.docx")
